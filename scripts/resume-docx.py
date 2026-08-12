"""Собирает docx из resume/<locale>.md — тот же источник, что и resume:pdf."""

import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Mm

SOURCE = Path(__file__).resolve().parent.parent / "resume"
OUT = Path(__file__).resolve().parent.parent / "resume" / "docx"

STYLES = {
    "classic": {
        "body_font": "Times New Roman",
        "display_font": "Times New Roman",
        "mono_font": "Times New Roman",
        "body_size": 10.5,
        "name_size": 21,
        "name_color": "1F3864",
        "name_caps": False,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "head_align": WD_ALIGN_PARAGRAPH.CENTER,
        "heading_color": "1F3864",
        "heading_size": 11.5,
        "heading_rule": ("bottom", "8FA8C8", 6),
        "heading_spacing": 0,
        "link_color": "2E5C8A",
        "accent_rule": None,
        "stack_mono": False,
    },
    "studio": {
        "body_font": "Arial",
        "display_font": "Arial",
        "mono_font": "Arial",
        "body_size": 10,
        "name_size": 22,
        "name_color": "000000",
        "name_caps": True,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "head_align": WD_ALIGN_PARAGRAPH.LEFT,
        "heading_color": "000000",
        "heading_size": 9,
        # Снизу, а не сверху: верхнюю рамку абзаца часть просмотрщиков не рисует
        "heading_rule": ("bottom", "000000", 18),
        "heading_spacing": 60,
        "link_color": "000000",
        "accent_rule": "BFA800",
        "stack_mono": True,
    },
}


def unescape(text):
    return re.sub(r"\\(.)", r"\1", text)


def strip_bold(block):
    return re.sub(r"^\*+|\*+$", "", block)


def parse(markdown):
    blocks = [b.strip() for b in re.split(r"\n\s*\n", markdown) if b.strip()]
    name, rest = blocks[0], blocks[1:]
    head, body = [], []
    for block in rest:
        if not body and not re.fullmatch(r"\*\*[^*]+\*\*", block):
            head.append(block)
        else:
            body.append(block)
    return strip_bold(name), head, body


INLINE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_hyperlink(paragraph, text, url, style):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), style["link_color"])
    props.append(color)
    link.append(run)
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    paragraph._p.append(link)


def write_inline(paragraph, text, style, **fmt):
    """Раскладывает markdown-ссылки в отдельные run'ы, остальное — обычным текстом."""
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            add_run(paragraph, unescape(text[position : match.start()]), style, **fmt)
        add_hyperlink(paragraph, unescape(match.group(1)), match.group(2), style)
        position = match.end()
    if position < len(text):
        add_run(paragraph, unescape(text[position:]), style, **fmt)


def add_run(paragraph, text, style, bold=False, italic=False, size=None, font=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font or style["body_font"]
    run.font.size = Pt(size or style["body_size"])
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


# Порядок детей `w:pPr` фиксирован схемой: `w:pBdr` идёт перед всем этим,
# иначе Word и Quick Look молча выбрасывают рамку
AFTER_BORDER = ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr")


def set_border(paragraph, edge, color, size):
    props = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)
    borders.append(border)
    props.insert_element_before(borders, *AFTER_BORDER)


def set_char_spacing(run, twentieths):
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(twentieths))
    run._element.get_or_add_rPr().append(spacing)


def paragraph(document, space_before=0, space_after=0, align=None, keep_next=False):
    p = document.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.15
    if align is not None:
        fmt.alignment = align
    fmt.keep_with_next = keep_next
    return p


def right_tab(p, document):
    section = document.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)


def render_block(document, block, style):
    heading = re.fullmatch(r"\*\*([^*]+)\*\*", block)
    if heading and heading.group(1) == heading.group(1).upper():
        p = paragraph(document, space_before=14, space_after=2, keep_next=True)
        run = add_run(
            p,
            heading.group(1),
            style,
            bold=True,
            size=style["heading_size"],
            font=style["display_font"],
            color=style["heading_color"],
        )
        if style["heading_spacing"]:
            set_char_spacing(run, style["heading_spacing"])
        edge, color, size = style["heading_rule"]
        set_border(p, edge, color, size)
        return

    if "\t" in block:
        title, period = block.split("\t", 1)
        p = paragraph(document, space_before=10, keep_next=True)
        right_tab(p, document)
        write_inline(p, strip_bold(title), style, bold=True, size=style["body_size"] + 0.5)
        add_run(p, "\t", style)
        write_inline(p, strip_bold(period.strip()), style, italic=True)
        return

    stack = re.fullmatch(r"\*\*\*(.+?)\*\*(.*)\*", block, re.S)
    if stack:
        p = paragraph(document, space_before=4)
        font = style["mono_font"] if style["stack_mono"] else style["body_font"]
        size = style["body_size"] - 1.5
        write_inline(p, stack.group(1), style, bold=True, italic=not style["stack_mono"], size=size, font=font)
        write_inline(p, stack.group(2), style, italic=not style["stack_mono"], size=size, font=font)
        return

    skill = re.fullmatch(r"\*\*(.+?:)\*\*\s*(.*)", block, re.S)
    if skill:
        p = paragraph(document, space_before=5)
        write_inline(p, skill.group(1), style, bold=True)
        add_run(p, " ", style)
        write_inline(p, skill.group(2), style)
        return

    subtitle = re.fullmatch(r"\*(.+)\*", block, re.S)
    if subtitle:
        p = paragraph(document, space_before=2, keep_next=True)
        write_inline(p, subtitle.group(1), style, italic=True, size=style["body_size"] - 0.5)
        return

    if block.startswith("* "):
        for item in re.split(r"\n(?=\* )", block):
            p = document.add_paragraph(style="List Bullet")
            fmt = p.paragraph_format
            fmt.space_before = Pt(2)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1.15
            fmt.left_indent = Mm(5)
            fmt.first_line_indent = Mm(-3.5)
            write_inline(p, re.sub(r"^\*\s+", "", item).strip(), style)
        return

    write_inline(paragraph(document, space_before=4), block, style)


def build(locale, style_name):
    style = STYLES[style_name]
    markdown = (SOURCE / f"{locale}.md").read_text(encoding="utf-8")
    name, head, body = parse(markdown)

    document = docx.Document()
    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.bottom_margin = Mm(16)
    section.left_margin = section.right_margin = Mm(21)

    normal = document.styles["Normal"]
    normal.font.name = style["body_font"]
    normal.font.size = Pt(style["body_size"])
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), style["body_font"])

    title = paragraph(document, align=style["name_align"], space_after=2)
    run = add_run(
        title,
        name.upper() if style["name_caps"] else name,
        style,
        bold=True,
        size=style["name_size"],
        font=style["display_font"],
        color=style["name_color"],
    )
    if style["name_caps"]:
        set_char_spacing(run, 40)
    if style["accent_rule"]:
        set_border(title, "bottom", style["accent_rule"], 24)

    for index, block in enumerate(head):
        p = paragraph(document, align=style["head_align"], space_before=3 if index else 2)
        write_inline(
            p,
            strip_bold(block),
            style,
            size=style["body_size"] + (3.5 if index == 0 else -0.5),
        )

    for block in body:
        render_block(document, block, style)

    OUT.mkdir(parents=True, exist_ok=True)
    target = OUT / f"Androsov_Viacheslav_Frontend_{locale.upper()}_{style_name}.docx"
    document.save(target)
    print(target)


if __name__ == "__main__":
    for style_name in sys.argv[1:] or STYLES:
        for locale in ("ru", "en"):
            build(locale, style_name)
